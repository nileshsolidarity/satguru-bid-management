"""
Satguru Bid Management — Flask server
Run: python3 server.py
Then open: http://localhost:3000
"""

import json
import os
import threading
import asyncio
from flask import Flask, send_from_directory, jsonify, request, Response
from flask_cors import CORS
from scraper import run_scraper

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
GCS_BUCKET = "satguru-bid-tenders"
GCS_OBJECT = "tenders.json"
LOCAL_TENDERS = os.path.join(BASE_DIR, "tenders.json")

app = Flask(__name__, static_folder=BASE_DIR)
CORS(app, origins=["https://satguru-bid-management.vercel.app", "http://localhost:3000"])

scraper_status = {"running": False, "last_run": None, "last_count": 0, "error": None}


def is_cloud():
    return os.environ.get("K_SERVICE") is not None  # set by Cloud Run automatically


def read_tenders():
    if is_cloud():
        try:
            from google.cloud import storage
            client = storage.Client()
            bucket = client.bucket(GCS_BUCKET)
            blob = bucket.blob(GCS_OBJECT)
            if blob.exists():
                return json.loads(blob.download_as_text())
        except Exception as e:
            print(f"GCS read error: {e}")
        return []
    else:
        if os.path.exists(LOCAL_TENDERS):
            with open(LOCAL_TENDERS) as f:
                return json.load(f)
        return []


def write_tenders(tenders):
    if is_cloud():
        try:
            from google.cloud import storage
            client = storage.Client()
            bucket = client.bucket(GCS_BUCKET)
            blob = bucket.blob(GCS_OBJECT)
            blob.upload_from_string(json.dumps(tenders, indent=2), content_type="application/json")
            print(f"Saved {len(tenders)} tenders to GCS")
        except Exception as e:
            print(f"GCS write error: {e}")
    else:
        with open(LOCAL_TENDERS, "w") as f:
            json.dump(tenders, f, indent=2)
        print(f"Saved {len(tenders)} tenders locally")


@app.route("/")
def index():
    return send_from_directory(BASE_DIR, "index.html")


@app.route("/<path:filename>")
def static_files(filename):
    return send_from_directory(BASE_DIR, filename)


@app.route("/api/sync", methods=["POST"])
def sync():
    if scraper_status["running"]:
        return jsonify({"status": "already_running", "message": "Scraper is already running..."}), 200

    def run():
        scraper_status["running"] = True
        scraper_status["error"] = None
        try:
            results = asyncio.run(run_scraper(headless=True, save_fn=write_tenders, load_fn=read_tenders))
            scraper_status["last_count"] = len(results)
            scraper_status["last_run"] = __import__("datetime").datetime.now().isoformat()
        except Exception as e:
            scraper_status["error"] = str(e)
        finally:
            scraper_status["running"] = False

    thread = threading.Thread(target=run, daemon=True)
    thread.start()
    return jsonify({"status": "started", "message": "Scraper started. Check back in 30-60 seconds."})


@app.route("/api/sync/status")
def sync_status():
    return jsonify(scraper_status)


@app.route("/api/tenders")
def get_tenders():
    return jsonify(read_tenders())


@app.route("/api/tender-detail")
def tender_detail():
    url = request.args.get("url", "").strip()
    if not url:
        return jsonify({"error": "url required"}), 400

    try:
        from scraper import fetch_tender_page
        html = asyncio.run(fetch_tender_page(url))
        return Response(html, mimetype="text/html")
    except Exception as e:
        return Response(f"<h2>Could not load tender</h2><p>{e}</p><p><a href='{url}' target='_blank'>Try opening directly</a></p>", mimetype="text/html")


@app.route("/api/debug-login")
def debug_login():
    """Debug endpoint: checks secret loading and tests GT login."""
    import asyncio
    from scraper import CREDS
    from playwright.async_api import async_playwright

    gt = CREDS["globaltenders"]
    result = {
        "gt_email": gt["email"],
        "gt_password_len": len(gt["password"]),
        "gt_password_set": bool(gt["password"]),
        "login_url": gt["login_url"],
    }

    async def test_login():
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            context = await browser.new_context(
                user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36"
            )
            page = await context.new_page()
            try:
                await page.goto(gt["login_url"], timeout=30000)
                await page.wait_for_timeout(2000)
                url_before = page.url

                # Check what form fields exist
                inputs = await page.evaluate("""
                    () => Array.from(document.querySelectorAll('input')).map(i => ({
                        name: i.name, type: i.type, id: i.id, placeholder: i.placeholder
                    }))
                """)

                if gt["password"]:
                    await page.fill('input[name="email"]', gt["email"])
                    await page.fill('input[name="password"]', gt["password"])
                    await page.click('input[type="submit"]')
                    await page.wait_for_timeout(5000)

                url_after = page.url
                page_text = await page.evaluate("() => document.body.innerText.slice(0, 500)")
                return {
                    "url_before_login": url_before,
                    "url_after_login": url_after,
                    "logged_in": "dashboard" in url_after.lower() or "logout" in page_text.lower(),
                    "page_snippet": page_text,
                    "form_inputs": inputs,
                }
            except Exception as e:
                return {"error": str(e)}
            finally:
                await browser.close()

    login_result = asyncio.run(test_login())
    result.update(login_result)
    return jsonify(result)


@app.route("/api/bid/generate", methods=["POST"])
def bid_generate():
    """Generate bid section content using Vertex AI Gemini 2.5 Flash."""
    body = request.get_json(force=True)
    section = body.get("section", "")
    context = body.get("context", {})  # tender details
    rfp_text = body.get("rfp_text", "")

    company_profile = """
Satguru Travel Group is a leading travel management company with over 30 years of experience.
Headquarters: Dubai, UAE. Offices across Africa, Middle East, and Asia.
Services: Corporate Travel Management, Air Ticketing, Hotel Bookings, Visa Services,
MICE (Meetings, Incentives, Conferences & Events), Car Rentals, Travel Insurance.
IATA accredited. ISO certified. Managed travel for government entities, NGOs, and Fortune 500 companies.
"""

    section_prompts = {
        "cover_letter": f"""Write a professional bid cover letter for a travel management tender.
Client: {context.get('client','')}. Tender Reference: {context.get('reference','')}. Bid Type: {context.get('bid_type','')}.
Use Satguru Travel Group as the bidder. Keep it formal, confident, 3 paragraphs. Do not use placeholders.""",

        "company_profile": f"""Write a company profile section for a bid proposal for {context.get('client','a client')}.
Use this factual information about the company: {company_profile}
Write 4-5 paragraphs covering: company overview, geographic presence, key services, experience, and why we are the right partner.
Professional tone, no placeholders.""",

        "understanding": f"""Write an 'Understanding of Requirements' section for a tender bid.
Client: {context.get('client','')}. Tender: {context.get('description','')}.
RFP Key Points: {rfp_text or 'Not provided - write a general understanding for a travel management bid'}.
Show deep understanding of the client's needs. 3-4 paragraphs. Professional tone.""",

        "methodology": f"""Write a 'Proposed Methodology / Approach' section for a travel management bid.
Client: {context.get('client','')}. Services required: {context.get('description','travel management services')}.
Cover: account management structure, technology platform, reporting, SLA commitments, escalation process.
4-5 paragraphs, professional and specific.""",

        "team": f"""Write a 'Key Personnel & Team Structure' section for a travel management bid for {context.get('client','')}.
Include: Account Manager role, Operations Team, 24/7 support desk, regional coordinators.
Describe the team structure and their responsibilities. 3 paragraphs. Professional tone. No actual names.""",

        "compliance": f"""Write a 'Compliance & Regulatory' section for a travel management bid.
Client: {context.get('client','')}. Cover: IATA accreditation, data protection / GDPR compliance,
financial stability, insurance coverage, certifications, and regulatory compliance in relevant markets.
2-3 paragraphs, factual and confident.""",

        "experience": f"""Write a 'Relevant Experience & Track Record' section for a travel management bid for {context.get('client','')}.
Highlight Satguru's experience with: government/NGO accounts, large corporate travel programs,
multi-country operations, volume handled. Mention types of clients (without naming specific companies).
3-4 paragraphs. Professional.""",
    }

    prompt = section_prompts.get(section)
    if not prompt:
        return jsonify({"error": f"Unknown section: {section}"}), 400

    try:
        import vertexai
        from vertexai.generative_models import GenerativeModel
        vertexai.init(project="satguru-sales-intel", location="us-central1")
        model = GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(prompt)
        return jsonify({"text": response.text})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/bid/export-docx", methods=["POST"])
def bid_export_docx():
    """Generate a formatted bid proposal Word document."""
    import io
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    body = request.get_json(force=True)
    doc_type = body.get("type", "technical")  # "technical" or "financial"
    context = body.get("context", {})
    sections = body.get("sections", {})  # section_name -> text content
    line_items = body.get("line_items", [])  # for financial

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    def add_heading(text, level=1, color=(26, 26, 46)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_para(text, bold=False, italic=False, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(8)
        return p

    def add_divider():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'E94560')
        pBdr.append(bottom)
        pPr.append(pBdr)

    client = context.get("client", "")
    reference = context.get("reference", "")
    bid_type = context.get("bid_type", "")
    today = __import__("datetime").date.today().strftime("%d %B %Y")

    if doc_type == "technical":
        # ── Cover Page ──
        doc.add_paragraph()
        doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("TECHNICAL PROPOSAL")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(26, 26, 46)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"Travel Management Services").bold = True

        doc.add_paragraph()
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f"Submitted to: {client}\n")
        info.add_run(f"Reference: {reference}\n")
        info.add_run(f"Bid Type: {bid_type}\n")
        info.add_run(f"Date: {today}")

        doc.add_paragraph()
        doc.add_paragraph()
        submitter = doc.add_paragraph()
        submitter.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = submitter.add_run("Satguru Travel Group")
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(233, 69, 96)
        submitter.add_run("\nwww.satgurutravel.com")

        doc.add_page_break()

        # ── Sections ──
        section_order = [
            ("cover_letter",    "Cover Letter"),
            ("company_profile", "1. Company Profile"),
            ("understanding",   "2. Understanding of Requirements"),
            ("methodology",     "3. Proposed Methodology & Approach"),
            ("experience",      "4. Relevant Experience"),
            ("team",            "5. Key Personnel & Team Structure"),
            ("compliance",      "6. Compliance & Regulatory"),
        ]
        for key, title in section_order:
            text = sections.get(key, "")
            if not text:
                continue
            add_heading(title, level=1)
            add_divider()
            doc.add_paragraph()
            for para in text.strip().split("\n\n"):
                if para.strip():
                    add_para(para.strip())
            doc.add_page_break()

    else:
        # ── Financial Proposal ──
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("FINANCIAL PROPOSAL")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(26, 26, 46)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"Travel Management Services — {client}")

        doc.add_paragraph()
        add_para(f"Client: {client}", bold=True)
        add_para(f"Tender Reference: {reference}")
        add_para(f"Date: {today}")
        add_para(f"Valid for: 90 days from date of submission")
        add_divider()
        doc.add_paragraph()

        add_heading("Pricing Summary", level=1)

        # Table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Service / Item", "Unit", "Rate (USD)", "Quantity", "Total (USD)"]):
            hdr[i].text = h
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1A1A2E')
            tcPr.append(shd)

        grand_total = 0
        for item in line_items:
            row = table.add_row().cells
            qty = float(item.get("qty", 1) or 1)
            rate = float(item.get("rate", 0) or 0)
            total = qty * rate
            grand_total += total
            row[0].text = item.get("service", "")
            row[1].text = item.get("unit", "")
            row[2].text = f"{rate:,.2f}"
            row[3].text = str(int(qty))
            row[4].text = f"{total:,.2f}"

        # Grand total row
        total_row = table.add_row().cells
        total_row[0].text = "GRAND TOTAL"
        total_row[3].text = ""
        total_row[4].text = f"USD {grand_total:,.2f}"
        for cell in total_row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        doc.add_paragraph()
        add_heading("Terms & Conditions", level=2)
        terms = [
            "All prices are quoted in USD unless otherwise specified.",
            "Prices are valid for 90 days from the date of submission.",
            "Payment terms: 30 days net from invoice date.",
            "Prices exclude applicable taxes and government levies unless stated.",
            "Service fees are subject to review annually upon contract renewal.",
            "Satguru Travel Group reserves the right to revise pricing based on fuel surcharge changes.",
        ]
        for t in terms:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(t).font.size = Pt(10)

        doc.add_paragraph()
        add_divider()
        sign = doc.add_paragraph()
        sign.add_run("\n\nAuthorised Signatory: _________________________\n\n")
        sign.add_run("Name: _________________________\n\n")
        sign.add_run("Designation: _________________________\n\n")
        sign.add_run("Date: _________________________\n\n")
        sign.add_run("Company Seal:")

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{'Technical' if doc_type == 'technical' else 'Financial'}_Proposal_{client.replace(' ','_')}_{today.replace(' ','_')}.docx"
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 3000))
    print(f"Satguru Bid Management running at http://localhost:{port}")
    app.run(host="0.0.0.0", port=port, debug=False)
